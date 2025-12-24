"""Tools for Discovery, Authoring, Quality, and Prioritization Agents

Implements all tool functions used by specialized agents in the requirements workflow.
"""

from typing import List, Dict, Optional, Any
from pydantic import BaseModel, Field
import re
from pathlib import Path

from .state import RequirementRaw, UserStory, QualityIssue, PrioritizedRequirement


# ============================================================================
# Discovery Agent Tools
# ============================================================================

class ExtractedRequirement(BaseModel):
    """Requirement extracted from a document."""
    title: str
    description: str
    type: str  # Functional | Non-Functional | Constraint
    source: str


class DocumentExtractionResult(BaseModel):
    """Result from document extraction."""
    requirements: List[ExtractedRequirement]
    metadata: Dict[str, Any] = Field(default_factory=dict)


def extract_from_document(
    file_path: str,
    file_type: str = "auto"
) -> DocumentExtractionResult:
    """Extract requirements from uploaded document (PDF, DOCX, TXT).
    
    Args:
        file_path: Path to the document file
        file_type: File type ("pdf", "docx", "txt", or "auto" to detect)
        
    Returns:
        DocumentExtractionResult with extracted requirements
        
    Raises:
        FileNotFoundError: If file doesn't exist
        ValueError: If file type not supported
    """
    path = Path(file_path)
    
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    # Auto-detect file type
    if file_type == "auto":
        file_type = path.suffix.lower().lstrip('.')
    
    # Extract text based on file type
    if file_type in ["txt", "md", "markdown"]:
        text = _extract_text_from_txt(path)
    elif file_type == "pdf":
        text = _extract_text_from_pdf(path)
    elif file_type in ["docx", "doc"]:
        text = _extract_text_from_docx(path)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")
    
    # Parse requirements from text
    requirements = _parse_requirements_from_text(text, source=path.name)
    
    return DocumentExtractionResult(
        requirements=requirements,
        metadata={
            "file_path": str(path),
            "file_type": file_type,
            "file_size": path.stat().st_size,
            "requirements_count": len(requirements)
        }
    )


def _extract_text_from_txt(path: Path) -> str:
    """Extract text from TXT file."""
    with open(path, 'r', encoding='utf-8', errors='ignore') as f:
        return f.read()


def _extract_text_from_pdf(path: Path) -> str:
    """Extract text from PDF file using pypdf."""
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ImportError("pypdf not installed. Install with: pip install pypdf")
    
    reader = PdfReader(path)
    text_parts = []
    
    for page in reader.pages:
        text_parts.append(page.extract_text())
    
    return "\n\n".join(text_parts)


def _extract_text_from_docx(path: Path) -> str:
    """Extract text from DOCX file using python-docx."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx not installed. Install with: pip install python-docx")
    
    doc = Document(path)
    text_parts = []
    
    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    
    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                text_parts.append(row_text)
    
    return "\n\n".join(text_parts)


def _parse_requirements_from_text(text: str, source: str) -> List[ExtractedRequirement]:
    """Parse requirements from extracted text.
    
    Looks for common patterns:
    - "The system shall/must/should..."
    - Numbered requirements (1., 2., etc.)
    - REQ-XXX: format
    - Bullet points with requirement-like content
    """
    requirements = []
    
    # Pattern 1: "The system shall/must/should..." 
    system_pattern = r'(?:the\s+system|application|platform|product)\s+(?:shall|must|should|will)\s+([^.]+\.)'
    matches = re.finditer(system_pattern, text, re.IGNORECASE)
    
    for match in matches:
        description = match.group(0).strip()
        title = _generate_title_from_description(description)
        
        requirements.append(ExtractedRequirement(
            title=title,
            description=description,
            type=_classify_requirement_type(description),
            source=source
        ))
    
    # Pattern 2: REQ-XXX format
    req_pattern = r'(REQ[-\s]?\d+)\s*[:\-]\s*([^\n]+)'
    matches = re.finditer(req_pattern, text, re.IGNORECASE)
    
    for match in matches:
        req_id = match.group(1)
        description = match.group(2).strip()
        
        requirements.append(ExtractedRequirement(
            title=f"{req_id}: {_generate_title_from_description(description)}",
            description=description,
            type=_classify_requirement_type(description),
            source=source
        ))
    
    # Pattern 3: Numbered list items that look like requirements
    numbered_pattern = r'^\s*\d+\.\s+([^\n]+)'
    for line in text.split('\n'):
        match = re.match(numbered_pattern, line)
        if match:
            description = match.group(1).strip()
            # Filter out non-requirement lines (TOC, headers, etc.)
            if len(description) > 20 and any(word in description.lower() for word in ['must', 'should', 'will', 'shall', 'can', 'able']):
                title = _generate_title_from_description(description)
                
                requirements.append(ExtractedRequirement(
                    title=title,
                    description=description,
                    type=_classify_requirement_type(description),
                    source=source
                ))
    
    # Deduplicate by description
    seen = set()
    unique_requirements = []
    for req in requirements:
        if req.description not in seen:
            seen.add(req.description)
            unique_requirements.append(req)
    
    return unique_requirements


def _generate_title_from_description(description: str, max_length: int = 60) -> str:
    """Generate a concise title from requirement description."""
    # Remove common prefixes
    title = re.sub(r'^(?:the\s+)?(?:system|application|platform|product)\s+(?:shall|must|should|will)\s+', '', description, flags=re.IGNORECASE)
    
    # Truncate to max length
    if len(title) > max_length:
        title = title[:max_length].rsplit(' ', 1)[0] + '...'
    
    # Capitalize first letter
    title = title[0].upper() + title[1:] if title else description[:max_length]
    
    return title.strip()


def _classify_requirement_type(description: str) -> str:
    """Classify requirement as Functional, Non-Functional, or Constraint."""
    description_lower = description.lower()
    
    # Non-functional indicators
    nfr_keywords = [
        'performance', 'speed', 'fast', 'scalable', 'available', 'reliability',
        'security', 'secure', 'usability', 'maintainable', 'portable',
        'response time', 'throughput', 'latency', 'uptime', 'audit'
    ]
    
    # Constraint indicators
    constraint_keywords = [
        'must use', 'must support', 'compatible with', 'integrate with',
        'comply with', 'standard', 'regulation', 'platform', 'technology'
    ]
    
    # Check for non-functional
    if any(keyword in description_lower for keyword in nfr_keywords):
        return "Non-Functional"
    
    # Check for constraints
    if any(keyword in description_lower for keyword in constraint_keywords):
        return "Constraint"
    
    # Default to functional
    return "Functional"


class RequirementValidationResult(BaseModel):
    """Result from requirement validation."""
    validation_passed: bool
    missing_fields: List[str] = Field(default_factory=list)
    suggestions: List[str] = Field(default_factory=list)


def validate_requirement_capture(requirement: RequirementRaw) -> RequirementValidationResult:
    """Validate that captured requirement has sufficient detail.
    
    Args:
        requirement: RequirementRaw object to validate
        
    Returns:
        RequirementValidationResult with validation status and feedback
    """
    missing_fields = []
    suggestions = []
    
    # Check required fields
    if not requirement.title or len(requirement.title) < 5:
        missing_fields.append("title")
        suggestions.append("Add a clear, concise title (at least 5 characters)")
    
    if not requirement.description or len(requirement.description) < 20:
        missing_fields.append("description")
        suggestions.append("Add a detailed description (at least 20 characters)")
    
    if not requirement.type:
        missing_fields.append("type")
        suggestions.append("Specify type: Functional, Non-Functional, Constraint, Assumption, or Risk")
    
    if not requirement.source:
        missing_fields.append("source")
        suggestions.append("Specify source: discovery_session, document, meeting, etc.")
    
    # Check for vague descriptions
    vague_terms = ['user-friendly', 'fast', 'efficient', 'good', 'bad', 'easy', 'simple']
    if any(term in requirement.description.lower() for term in vague_terms):
        suggestions.append("Description contains vague terms. Consider being more specific.")
    
    # Check if needs refinement flag
    if requirement.needs_refinement:
        suggestions.append("Requirement marked as needing refinement")
    
    validation_passed = len(missing_fields) == 0
    
    return RequirementValidationResult(
        validation_passed=validation_passed,
        missing_fields=missing_fields,
        suggestions=suggestions
    )


# ============================================================================
# Authoring Agent Tools
# ============================================================================

class UserStoryValidationResult(BaseModel):
    """Result from user story validation."""
    is_valid: bool
    improvement_suggestions: List[str] = Field(default_factory=list)
    invest_compliance: Dict[str, bool] = Field(default_factory=dict)


def validate_user_story(story: UserStory) -> UserStoryValidationResult:
    """Validate user story against INVEST principles and best practices.
    
    INVEST:
    - Independent: Can be developed separately
    - Negotiable: Details can be discussed
    - Valuable: Provides value to user
    - Estimable: Can be sized
    - Small: Fits in one iteration
    - Testable: Can be verified
    
    Args:
        story: UserStory object to validate
        
    Returns:
        UserStoryValidationResult with validation status and suggestions
    """
    suggestions = []
    invest_compliance = {
        "Independent": True,
        "Negotiable": True,
        "Valuable": False,
        "Estimable": False,
        "Small": True,
        "Testable": False
    }
    
    # Check story statement format: "As a [role], I want [feature] so that [benefit]"
    story_pattern = r'as\s+a\s+(\w+).*i\s+want.*so\s+that'
    if not re.search(story_pattern, story.story_statement.lower()):
        suggestions.append("Story statement should follow format: 'As a [role], I want [feature] so that [benefit]'")
        invest_compliance["Valuable"] = False
    else:
        # Check if benefit is present (Valuable)
        if 'so that' in story.story_statement.lower():
            invest_compliance["Valuable"] = True
        else:
            suggestions.append("Include 'so that [benefit]' to explain value")
    
    # Check acceptance criteria (Testable)
    if not story.acceptance_criteria or len(story.acceptance_criteria) == 0:
        suggestions.append("Add acceptance criteria to make story testable")
        invest_compliance["Testable"] = False
    else:
        invest_compliance["Testable"] = True
        
        # Check for vague acceptance criteria
        vague_criteria = [ac for ac in story.acceptance_criteria if len(ac) < 10]
        if vague_criteria:
            suggestions.append("Some acceptance criteria are too vague. Make them specific and measurable.")
    
    # Check effort estimate (Estimable)
    valid_estimates = ["XS", "S", "M", "L", "XL"]
    if story.effort_estimate not in valid_estimates:
        suggestions.append(f"Effort estimate must be one of: {', '.join(valid_estimates)}")
        invest_compliance["Estimable"] = False
    else:
        invest_compliance["Estimable"] = True
        
        # Warn if too large (Small)
        if story.effort_estimate in ["XL"]:
            suggestions.append("Consider breaking down XL stories into smaller stories")
            invest_compliance["Small"] = False
    
    # Check edge cases
    if not story.edge_cases or len(story.edge_cases) == 0:
        suggestions.append("Consider adding edge cases and error scenarios")
    
    # Check definition of done
    if not story.definition_of_done or len(story.definition_of_done) == 0:
        suggestions.append("Add definition of done checklist")
    
    is_valid = all(invest_compliance.values())
    
    return UserStoryValidationResult(
        is_valid=is_valid,
        improvement_suggestions=suggestions,
        invest_compliance=invest_compliance
    )


def format_user_story_template(
    role: str,
    feature: str,
    benefit: str
) -> str:
    """Format a user story in standard template.
    
    Args:
        role: User role (e.g., "customer", "admin")
        feature: Feature description
        benefit: Benefit/value to user
        
    Returns:
        Formatted user story statement
    """
    return f"As a {role}, I want {feature} so that {benefit}"


def validate_acceptance_criteria(criteria: List[str]) -> Dict[str, Any]:
    """Validate acceptance criteria for testability.
    
    Args:
        criteria: List of acceptance criteria
        
    Returns:
        Dict with validation results and suggestions
    """
    testable = []
    vague = []
    
    # Testable verbs
    testable_verbs = ['can', 'should', 'must', 'will', 'displays', 'shows', 'returns', 'validates', 'redirects']
    
    for criterion in criteria:
        criterion_lower = criterion.lower()
        
        # Check if contains testable verbs
        if any(verb in criterion_lower for verb in testable_verbs):
            testable.append(criterion)
        else:
            vague.append(criterion)
    
    return {
        "testable_count": len(testable),
        "vague_count": len(vague),
        "testable_criteria": testable,
        "vague_criteria": vague,
        "suggestions": [
            f"Make criterion more testable: {crit}" for crit in vague
        ] if vague else []
    }


# ============================================================================
# Quality Agent Tools
# ============================================================================

class QualityValidationResult(BaseModel):
    """Result from quality validation."""
    issues_found: List[QualityIssue] = Field(default_factory=list)
    total_issues: int = 0
    critical_count: int = 0
    high_count: int = 0
    medium_count: int = 0
    low_count: int = 0


def validate_requirements_quality(
    requirements: List[RequirementRaw],
    user_stories: Optional[List[UserStory]] = None
) -> QualityValidationResult:
    """Validate requirements quality across 4 dimensions.
    
    Dimensions:
    1. Ambiguity - Vague or unclear terms
    2. Completeness - Missing information
    3. Inconsistency - Contradictions
    4. Testability - Cannot be verified
    
    Args:
        requirements: List of requirements to validate
        user_stories: Optional list of user stories for cross-validation
        
    Returns:
        QualityValidationResult with all detected issues
    """
    issues = []
    issue_id_counter = 1
    
    # Dimension 1: Ambiguity Detection
    for req in requirements:
        ambiguity_issues = _detect_ambiguity(req, issue_id_counter)
        issues.extend(ambiguity_issues)
        issue_id_counter += len(ambiguity_issues)
    
    # Dimension 2: Completeness Check
    for req in requirements:
        completeness_issues = _check_completeness(req, issue_id_counter)
        issues.extend(completeness_issues)
        issue_id_counter += len(completeness_issues)
    
    # Dimension 3: Consistency Validation
    consistency_issues = _validate_consistency(requirements, issue_id_counter)
    issues.extend(consistency_issues)
    issue_id_counter += len(consistency_issues)
    
    # Dimension 4: Testability Check
    for req in requirements:
        testability_issues = _check_testability(req, user_stories, issue_id_counter)
        issues.extend(testability_issues)
        issue_id_counter += len(testability_issues)
    
    # Count by severity
    severity_counts = {
        "Critical": len([i for i in issues if i.severity == "Critical"]),
        "High": len([i for i in issues if i.severity == "High"]),
        "Medium": len([i for i in issues if i.severity == "Medium"]),
        "Low": len([i for i in issues if i.severity == "Low"])
    }
    
    return QualityValidationResult(
        issues_found=issues,
        total_issues=len(issues),
        critical_count=severity_counts["Critical"],
        high_count=severity_counts["High"],
        medium_count=severity_counts["Medium"],
        low_count=severity_counts["Low"]
    )


def _detect_ambiguity(req: RequirementRaw, start_id: int) -> List[QualityIssue]:
    """Detect ambiguous terms in requirement."""
    issues = []
    
    vague_terms = {
        'user-friendly': 'Define specific usability criteria (e.g., task completion time < 2 minutes)',
        'fast': 'Specify performance targets (e.g., response time < 200ms)',
        'efficient': 'Define efficiency metrics (e.g., CPU usage < 50%)',
        'scalable': 'Specify scaling requirements (e.g., support 10,000 concurrent users)',
        'reliable': 'Define reliability metrics (e.g., 99.9% uptime)',
        'secure': 'Specify security standards (e.g., OAuth 2.0, AES-256 encryption)',
        'easy to use': 'Define specific usability criteria',
        'good': 'Specify measurable quality criteria',
        'bad': 'Specify measurable quality criteria',
        'simple': 'Define what constitutes simplicity',
        'flexible': 'Specify what aspects should be configurable'
    }
    
    description_lower = req.description.lower()
    
    for term, fix in vague_terms.items():
        if term in description_lower:
            issues.append(QualityIssue(
                id=f"QA-{start_id:03d}",
                location=req.id,
                category="Ambiguity",
                severity="High" if term in ['secure', 'reliable', 'scalable'] else "Medium",
                description=f"Vague term '{term}' detected in requirement",
                recommended_fix=fix,
                status="Identified"
            ))
            start_id += 1
    
    return issues


def _check_completeness(req: RequirementRaw, start_id: int) -> List[QualityIssue]:
    """Check requirement completeness."""
    issues = []
    
    # Check description length
    if len(req.description) < 30:
        issues.append(QualityIssue(
            id=f"QA-{start_id:03d}",
            location=req.id,
            category="Incompleteness",
            severity="Medium",
            description="Requirement description is too brief",
            recommended_fix="Add more detail: who, what, why, when, where, how",
            status="Identified"
        ))
        start_id += 1
    
    # Check for missing context
    if req.type == "Functional" and not any(word in req.description.lower() for word in ['user', 'system', 'shall', 'must']):
        issues.append(QualityIssue(
            id=f"QA-{start_id:03d}",
            location=req.id,
            category="Incompleteness",
            severity="Low",
            description="Functional requirement lacks clear actor (user/system)",
            recommended_fix="Specify who performs the action and what the system does",
            status="Identified"
        ))
        start_id += 1
    
    return issues


def _validate_consistency(requirements: List[RequirementRaw], start_id: int) -> List[QualityIssue]:
    """Validate consistency across requirements."""
    issues = []
    
    # Check for duplicate titles
    titles = {}
    for req in requirements:
        if req.title in titles:
            issues.append(QualityIssue(
                id=f"QA-{start_id:03d}",
                location=f"{titles[req.title]}, {req.id}",
                category="Inconsistency",
                severity="Medium",
                description=f"Duplicate requirement title: '{req.title}'",
                recommended_fix="Merge duplicate requirements or differentiate titles",
                status="Identified"
            ))
            start_id += 1
        else:
            titles[req.title] = req.id
    
    # Check for contradictions (simple keyword-based)
    # This is a simplified check - real implementation would use NLP
    for i, req1 in enumerate(requirements):
        for req2 in requirements[i+1:]:
            if _are_potentially_conflicting(req1, req2):
                issues.append(QualityIssue(
                    id=f"QA-{start_id:03d}",
                    location=f"{req1.id}, {req2.id}",
                    category="Inconsistency",
                    severity="High",
                    description=f"Potential conflict detected between requirements",
                    recommended_fix="Review requirements for contradictions and resolve",
                    status="Identified"
                ))
                start_id += 1
    
    return issues


def _are_potentially_conflicting(req1: RequirementRaw, req2: RequirementRaw) -> bool:
    """Simple check for potential conflicts (keyword-based)."""
    # Look for negations in similar contexts
    words1 = set(req1.description.lower().split())
    words2 = set(req2.description.lower().split())
    
    overlap = words1 & words2
    
    # If significant overlap and one has negation, might be conflict
    if len(overlap) > 5:
        negations = {'not', 'no', 'never', 'cannot', 'must not', 'shall not'}
        has_negation_1 = bool(words1 & negations)
        has_negation_2 = bool(words2 & negations)
        
        return has_negation_1 != has_negation_2
    
    return False


def _check_testability(
    req: RequirementRaw,
    user_stories: Optional[List[UserStory]],
    start_id: int
) -> List[QualityIssue]:
    """Check if requirement is testable."""
    issues = []
    
    # Check if requirement has corresponding user story with acceptance criteria
    if user_stories:
        has_story = any(story.requirement_id == req.id for story in user_stories)
        
        if not has_story:
            issues.append(QualityIssue(
                id=f"QA-{start_id:03d}",
                location=req.id,
                category="Untestable",
                severity="Medium",
                description="No user story with acceptance criteria for this requirement",
                recommended_fix="Create user story with testable acceptance criteria",
                status="Identified"
            ))
            start_id += 1
    
    # Check for measurable criteria
    measurable_keywords = ['number', 'count', 'time', 'seconds', 'minutes', 'percent', '%', 'rate']
    if req.type == "Non-Functional" and not any(kw in req.description.lower() for kw in measurable_keywords):
        issues.append(QualityIssue(
            id=f"QA-{start_id:03d}",
            location=req.id,
            category="Untestable",
            severity="High",
            description="Non-functional requirement lacks measurable criteria",
            recommended_fix="Add specific metrics (e.g., response time < 200ms, uptime > 99.9%)",
            status="Identified"
        ))
        start_id += 1
    
    return issues


# ============================================================================
# Prioritization Agent Tools
# ============================================================================

class PrioritizationResult(BaseModel):
    """Result from prioritization framework application."""
    ranked_requirements: List[PrioritizedRequirement] = Field(default_factory=list)
    framework_used: str
    metadata: Dict[str, Any] = Field(default_factory=dict)


def apply_prioritization_framework(
    framework: str,
    requirements: List[RequirementRaw],
    scoring_inputs: Optional[Dict[str, Any]] = None
) -> PrioritizationResult:
    """Apply prioritization framework to rank requirements.
    
    Args:
        framework: RICE | MoSCoW | Kano | Value-Effort
        requirements: Requirements to prioritize
        scoring_inputs: Framework-specific inputs
        
    Returns:
        PrioritizationResult with ranked requirements
    """
    if framework == "RICE":
        return _apply_rice(requirements, scoring_inputs or {})
    elif framework == "MoSCoW":
        return _apply_moscow(requirements, scoring_inputs or {})
    elif framework == "Kano":
        return _apply_kano(requirements, scoring_inputs or {})
    elif framework == "Value-Effort":
        return _apply_value_effort(requirements, scoring_inputs or {})
    else:
        raise ValueError(f"Unknown framework: {framework}")


def _apply_rice(requirements: List[RequirementRaw], inputs: Dict[str, Any]) -> PrioritizationResult:
    """Apply RICE (Reach × Impact × Confidence / Effort) framework."""
    scored_requirements = []
    
    for i, req in enumerate(requirements):
        # Get or estimate RICE components (would be user-provided in real implementation)
        reach = inputs.get(req.id, {}).get('reach', 100)  # Number of users affected
        impact = inputs.get(req.id, {}).get('impact', 2)  # 0.25=Minimal, 0.5=Low, 1=Medium, 2=High, 3=Massive
        confidence = inputs.get(req.id, {}).get('confidence', 80)  # Percentage (0-100)
        effort = inputs.get(req.id, {}).get('effort', 5)  # Person-months
        
        # Calculate RICE score
        rice_score = (reach * impact * (confidence / 100)) / effort if effort > 0 else 0
        
        scored_requirements.append({
            'requirement': req,
            'score': rice_score
        })
    
    # Sort by score descending
    scored_requirements.sort(key=lambda x: x['score'], reverse=True)
    
    # Create prioritized requirements
    ranked = []
    for rank, item in enumerate(scored_requirements, start=1):
        req = item['requirement']
        score = item['score']
        
        # Assign priority level based on score quartiles
        if rank <= len(requirements) * 0.25:
            priority_level = "Must Have"
            phase = "Phase 1"
        elif rank <= len(requirements) * 0.5:
            priority_level = "Should Have"
            phase = "Phase 1"
        elif rank <= len(requirements) * 0.75:
            priority_level = "Could Have"
            phase = "Phase 2"
        else:
            priority_level = "Won't Have"
            phase = "Backlog"
        
        ranked.append(PrioritizedRequirement(
            rank=rank,
            requirement_id=req.id,
            title=req.title,
            priority_level=priority_level,
            framework_score=score,
            phase=phase,
            dependencies=[],
            enables=[],
            rationale=f"RICE score: {score:.2f} (Reach: {inputs.get(req.id, {}).get('reach', 100)}, "
                     f"Impact: {inputs.get(req.id, {}).get('impact', 2)}, "
                     f"Confidence: {inputs.get(req.id, {}).get('confidence', 80)}%, "
                     f"Effort: {inputs.get(req.id, {}).get('effort', 5)} PM)"
        ))
    
    return PrioritizationResult(
        ranked_requirements=ranked,
        framework_used="RICE",
        metadata={"total_requirements": len(requirements)}
    )


def _apply_moscow(requirements: List[RequirementRaw], inputs: Dict[str, Any]) -> PrioritizationResult:
    """Apply MoSCoW (Must, Should, Could, Won't) framework."""
    # Categorize requirements (would be user-driven in real implementation)
    categorized = {
        "Must Have": [],
        "Should Have": [],
        "Could Have": [],
        "Won't Have": []
    }
    
    for req in requirements:
        # Get category from inputs or default to Should Have
        category = inputs.get(req.id, {}).get('category', 'Should Have')
        categorized[category].append(req)
    
    # Create ranked list
    ranked = []
    rank = 1
    
    for priority_level in ["Must Have", "Should Have", "Could Have", "Won't Have"]:
        phase = "Phase 1" if priority_level in ["Must Have", "Should Have"] else "Phase 2" if priority_level == "Could Have" else "Backlog"
        
        for req in categorized[priority_level]:
            ranked.append(PrioritizedRequirement(
                rank=rank,
                requirement_id=req.id,
                title=req.title,
                priority_level=priority_level,
                framework_score=None,
                phase=phase,
                dependencies=[],
                enables=[],
                rationale=f"Categorized as {priority_level} based on business criticality"
            ))
            rank += 1
    
    return PrioritizationResult(
        ranked_requirements=ranked,
        framework_used="MoSCoW",
        metadata={"categories": {k: len(v) for k, v in categorized.items()}}
    )


def _apply_kano(requirements: List[RequirementRaw], inputs: Dict[str, Any]) -> PrioritizationResult:
    """Apply Kano model (Must-haves, Performance, Delighters)."""
    categorized = {
        "Must-haves": [],
        "Performance": [],
        "Delighters": []
    }
    
    for req in requirements:
        # Get category from inputs or classify based on type
        category = inputs.get(req.id, {}).get('category')
        if not category:
            # Default classification logic
            if req.type in ["Constraint", "Non-Functional"]:
                category = "Must-haves"
            elif "should" in req.description.lower():
                category = "Performance"
            else:
                category = "Delighters"
        
        categorized[category].append(req)
    
    # Prioritize: Must-haves > Performance > Delighters
    ranked = []
    rank = 1
    
    for category in ["Must-haves", "Performance", "Delighters"]:
        for req in categorized[category]:
            if category == "Must-haves":
                priority_level = "Must Have"
                phase = "Phase 1"
            elif category == "Performance":
                priority_level = "Should Have"
                phase = "Phase 1"
            else:
                priority_level = "Could Have"
                phase = "Phase 2"
            
            ranked.append(PrioritizedRequirement(
                rank=rank,
                requirement_id=req.id,
                title=req.title,
                priority_level=priority_level,
                framework_score=None,
                phase=phase,
                dependencies=[],
                enables=[],
                rationale=f"Kano category: {category} - {self._get_kano_rationale(category)}"
            ))
            rank += 1
    
    return PrioritizationResult(
        ranked_requirements=ranked,
        framework_used="Kano",
        metadata={"categories": {k: len(v) for k, v in categorized.items()}}
    )


def _get_kano_rationale(category: str) -> str:
    """Get rationale for Kano category."""
    rationales = {
        "Must-haves": "Basic expectations - absence causes dissatisfaction",
        "Performance": "Linear satisfaction - more is better",
        "Delighters": "Unexpected features - presence creates delight"
    }
    return rationales.get(category, "")


def _apply_value_effort(requirements: List[RequirementRaw], inputs: Dict[str, Any]) -> PrioritizationResult:
    """Apply Value-Effort matrix (2×2 grid)."""
    scored = []
    
    for req in requirements:
        # Get or estimate value and effort
        value = inputs.get(req.id, {}).get('value', 5)  # 1-10 scale
        effort = inputs.get(req.id, {}).get('effort', 5)  # 1-10 scale
        
        # Calculate simple score: value / effort
        score = value / effort if effort > 0 else 0
        
        # Categorize into quadrants
        if value >= 6 and effort <= 4:
            quadrant = "Quick Wins"
            priority = "Must Have"
            phase = "Phase 1"
        elif value >= 6 and effort > 4:
            quadrant = "Major Projects"
            priority = "Should Have"
            phase = "Phase 1"
        elif value < 6 and effort <= 4:
            quadrant = "Fill-ins"
            priority = "Could Have"
            phase = "Phase 2"
        else:
            quadrant = "Time Sinks"
            priority = "Won't Have"
            phase = "Backlog"
        
        scored.append({
            'requirement': req,
            'score': score,
            'quadrant': quadrant,
            'priority': priority,
            'phase': phase,
            'value': value,
            'effort': effort
        })
    
    # Sort by score descending
    scored.sort(key=lambda x: x['score'], reverse=True)
    
    # Create ranked list
    ranked = []
    for rank, item in enumerate(scored, start=1):
        req = item['requirement']
        
        ranked.append(PrioritizedRequirement(
            rank=rank,
            requirement_id=req.id,
            title=req.title,
            priority_level=item['priority'],
            framework_score=item['score'],
            phase=item['phase'],
            dependencies=[],
            enables=[],
            rationale=f"Value-Effort: {item['quadrant']} (Value: {item['value']}/10, Effort: {item['effort']}/10)"
        ))
    
    return PrioritizationResult(
        ranked_requirements=ranked,
        framework_used="Value-Effort",
        metadata={"total_requirements": len(requirements)}
    )


def analyze_dependencies(requirements: List[PrioritizedRequirement]) -> List[PrioritizedRequirement]:
    """Analyze and update dependency relationships between requirements.
    
    Args:
        requirements: List of prioritized requirements
        
    Returns:
        Updated requirements with dependencies populated
    """
    # This would use more sophisticated analysis in production
    # For now, simple keyword-based dependency detection
    
    for i, req in enumerate(requirements):
        title_lower = req.title.lower()
        
        # Look for dependencies
        for other_req in requirements:
            if other_req.requirement_id == req.requirement_id:
                continue
            
            other_title_lower = other_req.title.lower()
            
            # Simple keyword matching for dependencies
            if 'auth' in title_lower and 'login' in other_title_lower:
                if other_req.requirement_id not in req.dependencies:
                    req.dependencies.append(other_req.requirement_id)
            
            if 'user' in title_lower and 'authentication' in other_title_lower:
                if other_req.requirement_id not in req.dependencies:
                    req.dependencies.append(other_req.requirement_id)
    
    return requirements
